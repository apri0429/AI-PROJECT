from __future__ import annotations

import io
import logging
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path

import docx
import fitz
from PIL import Image
from pypdf import PdfReader, PdfWriter

from adapters.llm import (
    generate_json_from_image,
    translate_document_part,
    translate_pdf_document,
    translate_text,
)
from core.prompts import (
    build_diagram_box_refine_prompt,
    build_diagram_region_prompt,
    build_docx_translate_prompt,
    build_image_translate_prompt,
    build_pdf_translate_prompt,
)

# Extensions the translate endpoint accepts, and the mime type Gemini's file
# part needs for each image kind — anything outside PDF/DOCX is treated as
# a plain image (single "page") and sent to Gemini's vision path directly.
IMAGE_MIME_TYPES = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".bmp": "image/bmp",
    ".gif": "image/gif",
}
DOCX_EXTENSIONS = {".docx"}
SUPPORTED_EXTENSIONS = {".pdf", *IMAGE_MIME_TYPES, *DOCX_EXTENSIONS}

logger = logging.getLogger(__name__)

# Gemini reads each page directly (page image + text layer), which also lets
# it handle scanned/image-only PDFs without a separate OCR step. Translating
# one page per call (run concurrently) keeps each page's translated text
# paired with that exact page's diagram crops, so the two can be interleaved
# in the output doc in the original page order.
_MAX_CONCURRENT_PAGES = 4

# 2x zoom renders at roughly 144 DPI — sharp enough for Gemini to reliably
# spot diagram boundaries and for the cropped result to stay legible.
_RENDER_ZOOM = 2.0

# A detected box smaller than this fraction of the page's area is almost
# always a stray callout number or icon rather than an actual diagram.
_MIN_BOX_AREA_FRACTION = 0.01

# A small margin around each detected box so a diagram's outline/labels
# right at the box edge don't get clipped. Kept tight — too much padding
# starts dragging in stray fragments of neighboring caption text.
_BOX_PADDING_FRACTION = 0.004

# Margin used when re-cropping a box for the refine pass below — much
# looser than _BOX_PADDING_FRACTION since its purpose is the opposite:
# giving the refine call enough surrounding context to pull a clipped
# edge back in, not just avoid clipping labels sitting on the box edge.
_REFINE_MARGIN_FRACTION = 0.06


def _detect_diagram_boxes(page_image: Image.Image) -> list[dict]:
    """Ask Gemini vision to locate diagram/illustration regions on a
    rendered page (as opposed to running text), and return them as
    {"box": (left, top, right, bottom), "reference_marker": str | None}
    in page_image's own pixel coordinate space. Returns an empty list on
    any failure or if the page has no diagrams — callers should treat that
    as "nothing to crop", not an error."""
    buffer = io.BytesIO()
    page_image.save(buffer, format="PNG")

    try:
        raw_boxes = generate_json_from_image(build_diagram_region_prompt(), buffer.getvalue())
    except Exception:
        logger.exception("Diagram region detection failed")
        return []

    if not isinstance(raw_boxes, list):
        return []

    width, height = page_image.size
    pad_x, pad_y = int(width * _BOX_PADDING_FRACTION), int(height * _BOX_PADDING_FRACTION)
    min_area = width * height * _MIN_BOX_AREA_FRACTION

    boxes: list[dict] = []
    for item in raw_boxes:
        if not isinstance(item, dict):
            continue
        box_2d = item.get("box_2d")
        if not (isinstance(box_2d, list) and len(box_2d) == 4):
            continue
        try:
            ymin, xmin, ymax, xmax = (float(v) for v in box_2d)
        except (TypeError, ValueError):
            continue

        left = max(0, int(xmin / 1000 * width) - pad_x)
        top = max(0, int(ymin / 1000 * height) - pad_y)
        right = min(width, int(xmax / 1000 * width) + pad_x)
        bottom = min(height, int(ymax / 1000 * height) + pad_y)
        if right <= left or bottom <= top:
            continue
        if (right - left) * (bottom - top) < min_area:
            continue

        reference_marker = item.get("reference_marker")
        boxes.append({
            "box": (left, top, right, bottom),
            "reference_marker": str(reference_marker).strip() if reference_marker else None,
        })

    return boxes


def _refine_box(page_image: Image.Image, box: tuple[int, int, int, int]) -> tuple[int, int, int, int]:
    """The initial page-level detection in _detect_diagram_boxes works off a
    full page image, so its box edges are often a little off — clipping part
    of the drawing on one side, or dragging in a sliver of caption text on
    another. Re-crop with a generous margin around that box and ask Gemini
    to find the tight edge again, now zoomed into just this region where
    it's easier to place the boundary precisely. Falls back to the original
    box untouched on any failure, so a bad refine call never makes things
    worse than skipping it."""
    left, top, right, bottom = box
    width, height = page_image.size
    margin_x = max(4, int((right - left) * _REFINE_MARGIN_FRACTION))
    margin_y = max(4, int((bottom - top) * _REFINE_MARGIN_FRACTION))

    crop_left = max(0, left - margin_x)
    crop_top = max(0, top - margin_y)
    crop_right = min(width, right + margin_x)
    crop_bottom = min(height, bottom + margin_y)

    candidate = page_image.crop((crop_left, crop_top, crop_right, crop_bottom))
    crop_width, crop_height = candidate.size
    if crop_width <= 0 or crop_height <= 0:
        return box

    buffer = io.BytesIO()
    candidate.save(buffer, format="PNG")

    try:
        result = generate_json_from_image(build_diagram_box_refine_prompt(), buffer.getvalue())
    except Exception:
        logger.exception("Diagram box refine failed, keeping original box")
        return box

    if not isinstance(result, dict):
        return box
    box_2d = result.get("box_2d")
    if not (isinstance(box_2d, list) and len(box_2d) == 4):
        return box
    try:
        ymin, xmin, ymax, xmax = (float(v) for v in box_2d)
    except (TypeError, ValueError):
        return box

    refined_left = crop_left + int(xmin / 1000 * crop_width)
    refined_top = crop_top + int(ymin / 1000 * crop_height)
    refined_right = crop_left + int(xmax / 1000 * crop_width)
    refined_bottom = crop_top + int(ymax / 1000 * crop_height)

    if refined_right <= refined_left or refined_bottom <= refined_top:
        return box
    return (refined_left, refined_top, refined_right, refined_bottom)


def _extract_page_diagrams(page_image: Image.Image) -> list[dict]:
    """Crop out just the diagram/illustration regions of a rendered page
    (see _detect_diagram_boxes) as separate PNG images, in reading order
    (top to bottom, then left to right) — used instead of a screenshot of
    the whole page so the translated text isn't duplicated alongside it,
    and instead of raw embedded-image extraction since a scanned manual
    page is often saved as a single flattened image with no separate
    diagram/text layers to pull apart. Each result also carries the step/
    part marker (e.g. "NO.1") detected next to it, if any, so the caller
    can place it next to the matching numbered step in the translated text
    instead of just appending every diagram at the end."""
    boxes = _detect_diagram_boxes(page_image)
    boxes.sort(key=lambda item: (item["box"][1], item["box"][0]))

    crops: list[dict] = []
    for item in boxes:
        refined_box = _refine_box(page_image, item["box"])
        cropped = page_image.crop(refined_box)
        buffer = io.BytesIO()
        cropped.save(buffer, format="PNG")
        crops.append({
            "image_bytes": buffer.getvalue(),
            "reference_marker": item["reference_marker"],
        })
    return crops


def translate_pdf_to_indonesian(content: bytes) -> list[dict]:
    """Translate a PDF to Bahasa Indonesia page by page, keeping each page's
    translated text paired with that same page's diagram/illustration crops
    — so illustrations in the source PDF (e.g. an instruction manual's
    installation diagrams) can be re-inserted into the translated Google Doc
    alongside the translated text, instead of being discarded or duplicating
    the original English text. Returns a list of {"text": ..., "images": [...]}
    in original page order."""
    reader = PdfReader(io.BytesIO(content))
    num_pages = len(reader.pages)
    if num_pages == 0:
        raise ValueError("PDF has no pages")

    prompt = build_pdf_translate_prompt()

    def _translate_page(index: int) -> str:
        writer = PdfWriter()
        writer.add_page(reader.pages[index])
        buffer = io.BytesIO()
        writer.write(buffer)
        return translate_pdf_document(prompt, buffer.getvalue())

    pdf_doc = fitz.open(stream=content, filetype="pdf")
    try:
        matrix = fitz.Matrix(_RENDER_ZOOM, _RENDER_ZOOM)
        page_images = [
            Image.open(io.BytesIO(page.get_pixmap(matrix=matrix).tobytes("png")))
            for page in pdf_doc
        ]
    finally:
        pdf_doc.close()

    with ThreadPoolExecutor(max_workers=min(num_pages, _MAX_CONCURRENT_PAGES)) as executor:
        translated_texts = list(executor.map(_translate_page, range(num_pages)))
        page_diagrams = list(executor.map(_extract_page_diagrams, page_images))

    return [
        {"text": text, "images": images}
        for text, images in zip(translated_texts, page_diagrams)
    ]


def translate_image_to_indonesian(content: bytes, mime_type: str) -> list[dict]:
    """Translate a single photographed/scanned document page image to
    Bahasa Indonesia. Treated the same way a single PDF page is: Gemini
    reads the text directly off the image, and any diagrams/illustrations
    on it are cropped out and re-inserted next to the translated text.
    Returns a single-item list in the same {"text": ..., "images": [...]}
    shape as translate_pdf_to_indonesian, so both feed the same doc
    builder."""
    page_image = Image.open(io.BytesIO(content)).convert("RGB")
    text = translate_document_part(build_image_translate_prompt(), content, mime_type)
    images = _extract_page_diagrams(page_image)
    return [{"text": text, "images": images}]


def _extract_docx_text(document: docx.Document) -> str:
    """Pull paragraph and table cell text out of a docx in reading order.
    Gemini's document part support doesn't cover docx directly, so unlike
    the PDF/image paths this translates extracted plain text rather than
    reading the file itself — good enough for the text content, though it
    means embedded images below are attached as-is rather than positioned
    next to the paragraph they illustrate."""
    parts: list[str] = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _extract_docx_images(content: bytes) -> list[dict]:
    """Pull every embedded picture out of a docx's media parts — there's no
    per-paragraph position info here (see _extract_docx_text), so these are
    just returned in document order to be appended after the translated
    text, not interleaved with it."""
    document = docx.Document(io.BytesIO(content))
    images: list[dict] = []
    for rel in document.part.rels.values():
        if "image" not in rel.reltype:
            continue
        try:
            images.append({"image_bytes": rel.target_part.blob, "reference_marker": None})
        except Exception:
            logger.exception("Failed to read an embedded docx image, skipping it")
    return images


def translate_docx_to_indonesian(content: bytes) -> list[dict]:
    """Translate a Word (.docx) document to Bahasa Indonesia. Extracts the
    document's text and embedded images separately (see _extract_docx_text/
    _extract_docx_images) since Gemini's document-part reading doesn't cover
    docx the way it does PDF/images. Returns a single-item list in the same
    shape as translate_pdf_to_indonesian/translate_image_to_indonesian."""
    document = docx.Document(io.BytesIO(content))
    source_text = _extract_docx_text(document)
    if not source_text.strip():
        raise ValueError("Document has no readable text")

    text = translate_text(build_docx_translate_prompt(source_text))
    images = _extract_docx_images(content)
    return [{"text": text, "images": images}]


def translate_document_to_indonesian(content: bytes, filename: str) -> list[dict]:
    """Dispatch a translate-to-Indonesian request based on the uploaded
    file's extension — the single entry point the API route calls, so it
    doesn't need to know which of the PDF/image/docx code paths handled
    it."""
    extension = Path(filename or "").suffix.lower()
    if extension == ".pdf":
        return translate_pdf_to_indonesian(content)
    if extension in IMAGE_MIME_TYPES:
        return translate_image_to_indonesian(content, IMAGE_MIME_TYPES[extension])
    if extension in DOCX_EXTENSIONS:
        return translate_docx_to_indonesian(content)
    raise ValueError(
        "Unsupported file type. Upload a PDF, an image (JPG/PNG/WEBP/BMP/GIF), or a Word .docx file."
    )
